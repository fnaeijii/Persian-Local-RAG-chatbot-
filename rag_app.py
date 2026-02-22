import os
import re
import torch
import dill
import faiss
import numpy as np
import requests
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from sentence_transformers import SentenceTransformer, CrossEncoder
from transformers import AutoTokenizer
from rank_bm25 import BM25Okapi
import logging
import config
from typing import List, Tuple, Optional, Dict, Any
import warnings

warnings.filterwarnings('ignore')

# تنظیم logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)


# بررسی و تنظیم دستگاه (GPU/CPU)
def get_device():
    """تشخیص و انتخاب بهترین دستگاه موجود"""
    if torch.backends.mps.is_available():
        device = torch.device("mps")
        logger.info("🚀 Using Apple Silicon GPU (MPS)")
    elif torch.cuda.is_available():
        device = torch.device("cuda")
        logger.info("🚀 Using NVIDIA GPU (CUDA)")
    else:
        device = torch.device("cpu")
        logger.info("💻 Using CPU")
    return device


DEVICE = get_device()


# -----------------------
# Singleton Pattern برای مدل‌ها
# -----------------------
class ModelManager:
    """مدیریت Singleton برای بارگذاری یکباره مدل‌ها"""
    _instance = None
    _embedder = None
    _reranker = None
    _tokenizer = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super(ModelManager, cls).__new__(cls)
        return cls._instance

    def get_embedder(self):
        """بارگذاری یا بازگرداندن مدل embedding"""
        if self._embedder is None:
            logger.info("Loading embedding model...")
            try:
                self._embedder = SentenceTransformer(
                    config.EMBEDDING_MODEL,
                    device=DEVICE
                )
                logger.info(f"✅ Loaded embedding model: {config.EMBEDDING_MODEL} on {DEVICE}")
            except Exception as e:
                logger.warning(f"Failed to load {config.EMBEDDING_MODEL}: {e}")
                self._embedder = SentenceTransformer(
                    config.FALLBACK_EMBEDDING_MODEL,
                    device=DEVICE
                )
                logger.info(f"✅ Loaded fallback embedding model on {DEVICE}")
        return self._embedder

    def get_reranker(self):
        """بارگذاری یا بازگرداندن مدل reranker"""
        if self._reranker is None:
            logger.info("Loading reranker model...")
            try:
                self._reranker = CrossEncoder(
                    config.RERANKER_MODEL,
                    device=DEVICE,
                    max_length=512
                )
                logger.info(f"✅ Loaded reranker model: {config.RERANKER_MODEL} on {DEVICE}")
            except Exception as e:
                logger.warning(f"Failed to load {config.RERANKER_MODEL}: {e}")
                self._reranker = CrossEncoder(
                    config.FALLBACK_RERANKER_MODEL,
                    device=DEVICE,
                    max_length=512
                )
                logger.info(f"✅ Loaded fallback reranker model on {DEVICE}")
        return self._reranker

    def get_tokenizer(self):
        """بارگذاری یا بازگرداندن tokenizer"""
        if self._tokenizer is None:
            logger.info("Loading tokenizer...")
            self._tokenizer = AutoTokenizer.from_pretrained(config.TOKENIZER_MODEL)
            logger.info(f"✅ Loaded tokenizer: {config.TOKENIZER_MODEL}")
        return self._tokenizer


# ایجاد instance واحد از ModelManager
model_manager = ModelManager()



def normalize_persian(text: str) -> str:
    """نرمال‌سازی متن فارسی"""
    if not text:
        return ""

    # تبدیل ی و ک عربی به فارسی
    text = text.replace("ي", "ی").replace("ك", "ک")

    # حذف فاصله‌های اضافی
    text = re.sub(r"\s+", " ", text)

    # حذف خط تیره‌های اضافی
    text = re.sub(r"ـ+", "", text)

    # حذف فاصله‌های نیم‌فاصله زائد
    text = text.replace("\u200c", " ")
    text = re.sub(r"\s+", " ", text)

    return text.strip()


def tok_simple_fa(s: str) -> List[str]:
    """توکنایزر ساده برای متن فارسی"""
    s = normalize_persian(s)
    s = re.sub(r"[^\w\s]", " ", s)
    return [w for w in s.split() if w]


def dedup_keep_order(items: List) -> List:
    """حذف تکراری‌ها با حفظ ترتیب"""
    seen, out = set(), []
    for x in items:
        if x not in seen:
            out.append(x)
            seen.add(x)
    return out


class OptimizedE5Embedder:
    """کلاس بهینه‌شده برای embedding با پشتیبانی batch processing"""

    def __init__(self):
        self.model = model_manager.get_embedder()
        self.batch_size = 32 if DEVICE.type != 'cpu' else 8

    def embed_passages(self, texts: List[str]) -> np.ndarray:
        """تولید embedding برای متن‌ها با batch processing"""
        texts = [f"passage: {t}" for t in texts]

        # Batch processing برای سرعت بیشتر
        with torch.no_grad():
            vecs = self.model.encode(
                texts,
                convert_to_numpy=True,
                show_progress_bar=True,
                batch_size=self.batch_size,
                normalize_embeddings=True
            )

        return vecs

    def embed_query(self, q: str) -> np.ndarray:
        """تولید embedding برای query"""
        with torch.no_grad():
            vec = self.model.encode(
                [f"query: {q}"],
                convert_to_numpy=True,
                normalize_embeddings=True
            )
        return vec


class OptimizedReranker:
    """کلاس بهینه‌شده برای reranking"""

    def __init__(self):
        self.model = model_manager.get_reranker()
        self.batch_size = 16 if DEVICE.type != 'cpu' else 4

    def rerank(self, query: str, candidates: List[str], top_k: int = config.RERANK_TOP_K) -> List[Tuple[str, float]]:
        """rerank کردن نامزدها"""
        if not candidates:
            return []

        pairs = [(query, c) for c in candidates]

        # Batch prediction برای سرعت بیشتر
        with torch.no_grad():
            scores = self.model.predict(
                pairs,
                batch_size=self.batch_size,
                show_progress_bar=False
            )

        return sorted(zip(candidates, scores), key=lambda x: x[1], reverse=True)[:top_k]


class TextFileLoader:
    def __init__(self, txt_path: str):
        self.txt_path = txt_path

    def load_text(self) -> str:
        """بارگذاری و پردازش متن از فایل TXT"""
        if not os.path.exists(self.txt_path):
            raise FileNotFoundError(f"Text file not found: {self.txt_path}")

        # تشخیص encoding فایل
        content = None
        for encoding in config.SUPPORTED_ENCODINGS:
            try:
                with open(self.txt_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    logger.info(f"Successfully read file with encoding: {encoding}")
                    break
            except UnicodeDecodeError:
                continue

        if content is None:
            with open(self.txt_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                logger.warning("File read with encoding errors ignored")

        # نرمال‌سازی متن
        content = normalize_persian(content)

        # پردازش خطوط و ساخت پاراگراف‌ها
        lines = content.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)

        # بازسازی متن با پاراگراف‌های مناسب
        paragraphs = []
        current_para = []

        for line in cleaned_lines:
            # تشخیص پاراگراف جدید
            if (re.match(r'^\d+[\.\-\)]', line) or
                    re.match(r'^[\*\-\•\◦]', line) or
                    len(line) < 50):

                if current_para:
                    paragraphs.append(' '.join(current_para))
                    current_para = []
                paragraphs.append(line)
            else:
                current_para.append(line)

        # آخرین پاراگراف
        if current_para:
            paragraphs.append(' '.join(current_para))

        # ترکیب پاراگراف‌ها
        full_text = '\n\n'.join(paragraphs)
        return full_text


class ImprovedChunker:
    def __init__(self, max_tokens: int = config.MAX_TOKENS, overlap: int = config.OVERLAP_TOKENS):
        self.tokenizer = model_manager.get_tokenizer()
        self.max_tokens = max_tokens
        self.overlap = overlap

    def chunk_by_paragraphs(self, text: str) -> List[str]:
        """چانک‌سازی هوشمند با حفظ ساختار و معنا"""

        # الگوهای تشخیص عناوین و سرفصل‌ها
        title_patterns = [
            r'^#+\s+',  # Markdown headers
            r'^\d+[\.\-\)]\s+',  # شماره‌گذاری
            r'^[•·▪▫◦‣⁃]\s+',  # بولت پوینت‌ها
            r'^فصل\s+\d+',  # الگوی فصل
            r'^بخش\s+\d+',  # الگوی بخش
        ]

        # تقسیم متن به پاراگراف‌ها
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

        chunks = []
        current_chunk = ""
        current_tokens = 0

        for para in paragraphs:
            # بررسی آیا پاراگراف یک عنوان است
            is_title = any(re.match(pattern, para) for pattern in title_patterns)

            # محاسبه تعداد توکن‌های پاراگراف
            para_tokens = len(self.tokenizer(para, add_special_tokens=False)["input_ids"])

            # اگر عنوان است و چانک فعلی خالی نیست، چانک جدید شروع کن
            if is_title and current_chunk:
                chunks.append(normalize_persian(current_chunk))
                # Overlap: حفظ قسمتی از چانک قبلی
                overlap_lines = current_chunk.split('\n')[-2:] if '\n' in current_chunk else []
                overlap_text = '\n'.join(overlap_lines) if overlap_lines else ''
                current_chunk = overlap_text + '\n\n' + para if overlap_text else para
                current_tokens = len(self.tokenizer(current_chunk, add_special_tokens=False)["input_ids"])

            # اگر پاراگراف خیلی بزرگ است
            elif para_tokens > self.max_tokens:
                # ذخیره چانک فعلی
                if current_chunk:
                    chunks.append(normalize_persian(current_chunk))
                    current_chunk = ""
                    current_tokens = 0

                # تقسیم پاراگراف بزرگ
                sentences = re.split(r'[.!?؟۔।।॥]', para)
                temp_chunk = ""
                temp_tokens = 0

                for sent in sentences:
                    sent = sent.strip()
                    if not sent:
                        continue

                    sent_tokens = len(self.tokenizer(sent, add_special_tokens=False)["input_ids"])

                    if temp_tokens + sent_tokens > self.max_tokens:
                        if temp_chunk:
                            chunks.append(normalize_persian(temp_chunk))
                        temp_chunk = sent + "."
                        temp_tokens = sent_tokens
                    else:
                        temp_chunk += " " + sent + "." if temp_chunk else sent + "."
                        temp_tokens += sent_tokens

                if temp_chunk:
                    chunks.append(normalize_persian(temp_chunk))

            # اگر اضافه کردن پاراگراف باعث عبور از حد مجاز شود
            elif current_tokens + para_tokens > self.max_tokens:
                if current_chunk:
                    chunks.append(normalize_persian(current_chunk))
                current_chunk = para
                current_tokens = para_tokens

            # در غیر این صورت پاراگراف را اضافه کن
            else:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
                current_tokens += para_tokens

        # آخرین چانک
        if current_chunk:
            chunks.append(normalize_persian(current_chunk))

        # فیلتر کردن چانک‌های کوتاه
        return [chunk for chunk in chunks if len(chunk) >= config.MIN_CHUNK_LENGTH]


class IndexCache:
    """مدیریت cache برای ایندکس‌ها"""
    _cache: Dict[str, Any] = {}

    @classmethod
    def get(cls, key: str) -> Optional[Any]:
        return cls._cache.get(key)

    @classmethod
    def set(cls, key: str, value: Any):
        cls._cache[key] = value

    @classmethod
    def clear(cls):
        cls._cache.clear()


def save_index(index, paragraphs: List[str], embeddings: np.ndarray,
               index_bin_path: str = config.INDEX_BIN_PATH,
               meta_path: str = config.META_PKL_PATH):
    """ذخیره ایندکس و متادیتا"""
    faiss.write_index(index, index_bin_path)
    with open(meta_path, "wb") as f:
        dill.dump((paragraphs, embeddings), f)

    # ذخیره در cache
    IndexCache.set('index', index)
    IndexCache.set('paragraphs', paragraphs)
    IndexCache.set('embeddings', embeddings)

    logger.info(f"Index saved to {index_bin_path} and {meta_path}")


def load_index(index_bin_path: str = config.INDEX_BIN_PATH,
               meta_path: str = config.META_PKL_PATH) -> Tuple:
    """بارگذاری ایندکس با استفاده از cache"""
    # بررسی cache
    cached_index = IndexCache.get('index')
    cached_paragraphs = IndexCache.get('paragraphs')
    cached_embeddings = IndexCache.get('embeddings')

    if cached_index is not None and cached_paragraphs is not None and cached_embeddings is not None:
        logger.info("Loading index from cache")
        return cached_index, cached_paragraphs, cached_embeddings

    # بارگذاری از فایل
    index = faiss.read_index(index_bin_path)
    with open(meta_path, "rb") as f:
        paragraphs, embeddings = dill.load(f)

    # ذخیره در cache
    IndexCache.set('index', index)
    IndexCache.set('paragraphs', paragraphs)
    IndexCache.set('embeddings', embeddings)

    logger.info(f"Index loaded from {index_bin_path} and {meta_path}")
    return index, paragraphs, embeddings



def build_index(txt_path: str, index_bin: str = config.INDEX_BIN_PATH,
                meta_pkl: str = config.META_PKL_PATH) -> str:
    """ساخت ایندکس از فایل متنی و بازگرداندن محتوای فایل"""
    print("📥 در حال ساخت ایندکس از فایل متنی...")
    logger.info(f"Building index from: {txt_path}")

    # بارگذاری متن
    loader = TextFileLoader(txt_path)
    text = loader.load_text()

    print(f"📖 متن با {len(text)} کاراکتر بارگذاری شد.")
    logger.info(f"Text loaded: {len(text)} characters")

    # چانک‌سازی با حفظ ساختار
    chunker = ImprovedChunker()
    paragraphs = chunker.chunk_by_paragraphs(text)

    # حذف چانک‌های کوتاه و تکراری
    paragraphs = dedup_keep_order([c for c in paragraphs if len(c) >= config.MIN_CHUNK_LENGTH])

    print(f"📄 {len(paragraphs)} چانک با کیفیت ساخته شد.")
    logger.info(f"Created {len(paragraphs)} quality chunks")

    # نمایش نمونه چانک‌ها
    print("\n📋 نمونه چانک‌های ایجاد شده:")
    for i, chunk in enumerate(paragraphs[:3]):
        print(f"\nچانک {i + 1}:")
        print(chunk[:200] + "..." if len(chunk) > 200 else chunk)
        print("-" * 50)

    # تولید embedding ها
    embedder = OptimizedE5Embedder()
    embeddings = embedder.embed_passages(paragraphs)

    # ساخت ایندکس FAISS
    dim = embeddings.shape[1]

    # استفاده از ایندکس بهینه برای GPU
    if DEVICE.type != 'cpu':
        # برای GPU از IndexFlatIP استفاده می‌کنیم
        faiss_index = faiss.IndexFlatIP(dim)
    else:
        # برای CPU می‌توانیم از IndexHNSWFlat استفاده کنیم
        faiss_index = faiss.IndexHNSWFlat(dim, 32)
        faiss_index.hnsw.efConstruction = 40

    faiss_index.add(embeddings)

    # ذخیره ایندکس
    save_index(faiss_index, paragraphs, embeddings, index_bin, meta_pkl)
    print("✅ ایندکس ذخیره شد.")

    # بازگرداندن محتوای فایل برای نمایش در GUI
    return text



def rrf_fuse(dense_ranked: List[int], bm25_ranked: List[int],
             k: int = config.RRF_K, topn: int = config.RRF_TOP_N) -> List[int]:
    """Reciprocal Rank Fusion برای ترکیب نتایج"""
    from collections import defaultdict
    score = defaultdict(float)

    for r, idx in enumerate(dense_ranked, 1):
        score[idx] += 1.0 / (k + r)
    for r, idx in enumerate(bm25_ranked, 1):
        score[idx] += 1.0 / (k + r)

    return [i for i, _ in sorted(score.items(), key=lambda x: x[1], reverse=True)[:topn]]


def expand_query(query: str) -> List[str]:
    """گسترش query فارسی برای جستجوی بهتر"""

    queries = [query]

    # حذف کلمات اضافی
    stop_words = ['که', 'را', 'از', 'به', 'در', 'با', 'برای', 'این', 'آن']
    clean_query = ' '.join([w for w in query.split() if w not in stop_words])
    if clean_query != query and len(clean_query) > 3:
        queries.append(clean_query)

    return queries[:2]  # حداکثر 2 query


def search_by_threshold(query: str,
                threshold: float = None,
                index_bin: str = config.INDEX_BIN_PATH,
                meta_pkl: str = config.META_PKL_PATH) -> List[Tuple[str, float]]:
    """جستجوی بهینه‌شده پاراگراف‌های مرتبط با امتیاز بالای threshold"""

    # استفاده از threshold پیش‌فرض از config اگر مقدار داده نشده
    if threshold is None:
        threshold = config.SCORE_THRESHOLD

    logger.info(f"Searching for query: {query} with threshold: {threshold}")

    # گسترش query برای جستجوی بهتر
    expanded_queries = expand_query(query)

    # بارگذاری ایندکس و مدل‌ها
    embedder = OptimizedE5Embedder()
    reranker = OptimizedReranker()
    faiss_index, paragraphs, embeddings = load_index(index_bin, meta_pkl)

    # ساخت BM25 index
    bm25 = BM25Okapi([tok_simple_fa(p) for p in paragraphs])

    # Dense retrieval
    q_emb = embedder.embed_query(query)
    distances, indices = faiss_index.search(q_emb, config.DENSE_SEARCH_TOP_K)
    dense_order = list(map(int, indices[0]))

    # BM25 retrieval
    scores = bm25.get_scores(tok_simple_fa(query))
    bm25_order = list(np.argsort(scores)[::-1][:config.BM25_SEARCH_TOP_K])

    # Fusion
    cand_order = rrf_fuse(dense_order, bm25_order)
    candidates = [paragraphs[i] for i in cand_order]

    # Reranking - ابتدا با تعداد بیشتری کاندید برای داشتن گزینه‌های کافی
    initial_rerank_k = min(len(candidates), config.MAX_PARAGRAPHS * 2)
    reranked_results = reranker.rerank(query, candidates, top_k=initial_rerank_k)

    # فیلتر کردن نتایج بر اساس threshold
    filtered_results = []
    for para, score in reranked_results:
        if score >= threshold:
            filtered_results.append((para, score))
            # اگر به حداکثر تعداد رسیدیم، متوقف شو
            if len(filtered_results) >= config.MAX_PARAGRAPHS:
                break

    # اگر هیچ نتیجه‌ای بالای threshold نبود، حداقل تعداد را برگردان
    if len(filtered_results) == 0 and config.MIN_PARAGRAPHS > 0:
        logger.warning(f"No results above threshold {threshold}. Returning top {config.MIN_PARAGRAPHS} results.")
        filtered_results = reranked_results[:config.MIN_PARAGRAPHS]

    # اطمینان از اینکه تعداد نتایج از حداکثر بیشتر نشود
    if len(filtered_results) > config.MAX_PARAGRAPHS:
        filtered_results = filtered_results[:config.MAX_PARAGRAPHS]

    logger.info(
        f"Found {len(filtered_results)} relevant paragraphs (threshold: {threshold}, scores: {[f'{score:.3f}' for _, score in filtered_results]})")

    return filtered_results

def strip_think_sections(text: str) -> str:
    """حذف کامل بخش‌های thinking از پاسخ"""
    if not text:
        return ""

    # حذف تگ‌های مختلف thinking
    patterns = [
        r'<think>.*?</think>',
        r'<thinking>.*?</thinking>',
        r'<thought>.*?</thought>',
        r'<\|thinking\|>.*?<\|/thinking\|>',
    ]

    for pattern in patterns:
        text = re.sub(pattern, '', text, flags=re.DOTALL | re.IGNORECASE)

    # حذف خطوط خالی اضافی
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def extract_ollama_response(response_data: Dict[str, Any]) -> str:
    """استخراج پاسخ نهایی از response object Ollama - نسخه بهبود یافته"""
    try:
        # اول بررسی می‌کنیم که آیا response کامل شده است
        if response_data.get('done', False):

            # روش 1: دریافت از message.content (برای مدل‌های جدید)
            if 'message' in response_data:
                message = response_data['message']
                if isinstance(message, dict):
                    # حذف بخش thinking اگر وجود داشته باشد
                    thinking = message.get('thinking', '')
                    content = message.get('content', '')

                    # اگر content خالی است اما thinking پر است
                    if not content and thinking:
                        # استخراج پاسخ از thinking
                        logger.warning("Content is empty, extracting from thinking section")
                        # جستجو برای پاسخ واقعی در thinking
                        if 'جواب' in thinking or 'پاسخ' in thinking:
                            # سعی در استخراج بخش پاسخ
                            lines = thinking.split('\n')
                            answer_started = False
                            answer_lines = []

                            for line in lines:
                                if any(word in line for word in ['جواب', 'پاسخ', 'بنابراین']):
                                    answer_started = True
                                if answer_started:
                                    answer_lines.append(line)

                            if answer_lines:
                                content = '\n'.join(answer_lines)

                        # اگر هنوز content خالی است، کل thinking را برگردان
                        if not content:
                            content = thinking

                    # پاکسازی content از تگ‌های thinking
                    if content:
                        content = strip_think_sections(content)
                        # اگر content فقط فضای خالی است
                        if content.strip():
                            return content.strip()

            # روش 2: دریافت مستقیم از response
            if 'response' in response_data:
                content = response_data.get('response', '')
                if content:
                    return strip_think_sections(content)

            # روش 3: دریافت از content مستقیم
            if 'content' in response_data:
                content = response_data.get('content', '')
                if content:
                    return strip_think_sections(content)

        # اگر هیچ‌کدام کار نکرد
        logger.error(f"Could not extract proper response. Keys available: {response_data.keys()}")

        # آخرین تلاش: بررسی دستی message
        if 'message' in response_data:
            msg = response_data['message']
            logger.warning(f"Message structure: {msg.keys() if isinstance(msg, dict) else type(msg)}")

        return "متأسفانه نتوانستم پاسخ را از مدل استخراج کنم. لطفاً دوباره تلاش کنید."

    except Exception as e:
        logger.error(f"Error extracting Ollama response: {e}")
        logger.error(f"Response structure: {json.dumps(response_data, indent=2, ensure_ascii=False)[:500]}")
        return "خطا در استخراج پاسخ از مدل"



def post_process_answer(answer: str) -> str:
    """پردازش نهایی پاسخ برای بهبود کیفیت و فرمت نمایش"""

    # حذف "پاسخ:" از ابتدای متن اگر وجود داشته باشد
    if answer.startswith("پاسخ:"):
        answer = answer[5:].strip()

    # تبدیل markdown به فرمت ساده
    # حذف ** برای bold
    answer = re.sub(r'\*\*(.*?)\*\*', r'\1', answer)

    # حذف * تکی برای لیست‌ها
    answer = re.sub(r'^\s*\*\s+', '• ', answer, flags=re.MULTILINE)
    answer = re.sub(r'\n\s*\*\s+', '\n• ', answer)

    # تبدیل شماره‌گذاری به فرمت بهتر
    answer = re.sub(r'^(\d+)\.\s+\*\*', r'\1. ', answer, flags=re.MULTILINE)
    answer = re.sub(r'\n(\d+)\.\s+\*\*', r'\n\1. ', answer)

    # اضافه کردن خط جدید بعد از نقطه‌های پایان جمله (برای خوانایی بهتر)
    # ابتدا جملات رو جدا می‌کنیم
    sentences = []
    current = ""

    for char in answer:
        current += char
        # اگر به پایان جمله رسیدیم
        if char in '.؟!':
            # بررسی که این نقطه جزء عدد نباشد (مثل 1.)
            if not (len(current) > 1 and current[-2].isdigit()):
                sentences.append(current.strip())
                current = ""

    # اگر چیزی باقی مانده
    if current.strip():
        sentences.append(current.strip())

    # بازسازی متن با فرمت بهتر
    formatted_parts = []
    for sent in sentences:
        sent = sent.strip()
        if not sent:
            continue

        # اگر جمله با شماره شروع می‌شود (لیست شماره‌دار)
        if re.match(r'^\d+\.', sent):
            formatted_parts.append('\n' + sent)
        # اگر جمله با bullet point شروع می‌شود
        elif sent.startswith('•'):
            formatted_parts.append('\n  ' + sent)
        # جملات عادی
        else:
            formatted_parts.append(sent)

    # ترکیب قسمت‌ها
    answer = '\n'.join(formatted_parts)

    # پاکسازی نهایی
    # حذف خطوط خالی اضافی
    answer = re.sub(r'\n{3,}', '\n\n', answer)

    # حذف فضاهای اضافی
    answer = re.sub(r'[ \t]+', ' ', answer)

    # حذف فضای ابتدا و انتهای خطوط
    lines = [line.strip() for line in answer.split('\n')]
    answer = '\n'.join(lines)

    # اطمینان از نقطه پایانی
    if answer and answer[-1] not in '.!?؟':
        answer += '.'

    return answer.strip()


def ask_deepseek(paragraphs: List[str], question: str,
                 model: str = config.OLLAMA_MODEL,
                 ollama_url: str = config.OLLAMA_URL) -> str:
    """ارسال سوال به Ollama با پرامپت بهینه‌شده"""

    # ساخت context از پاراگراف‌ها
    context = "\n\n".join(paragraphs)

    # پرامپت بهینه‌شده
    prompt = f"""<|system|>
شما یک دستیار هوش مصنوعی متخصص در پاسخگویی دقیق به سوالات فارسی هستید.

قوانین پاسخگویی:
1. فقط و فقط بر اساس متن ارائه شده پاسخ دهید
2. اگر اطلاعات کافی در متن نیست، صراحتاً اعلام کنید
3. پاسخ‌ها باید دقیق، مختصر و مفید باشند
4. از حدس و گمان پرهیز کنید
5. ساختار پاسخ باید منطقی و منسجم باشد
<|user|>

📚 متن مرجع:
================
{context}
================

❓ سوال: {question}

لطفاً با توجه به متن بالا، پاسخی جامع و دقیق ارائه دهید.
<|assistant|>
"""

    payload = {
        "model": model,
        "messages": [
            {'role': 'user', 'content': prompt}
        ],
        "options": config.OLLAMA_ADVANCED_OPTIONS,
        "stream": False,
    }

    max_retries = 3
    retry_count = 0

    while retry_count < max_retries:
        try:
            logger.info(f"Sending request to Ollama (attempt {retry_count + 1}/{max_retries})...")

            resp = requests.post(
                ollama_url.rstrip("/") + "/api/chat",
                json=payload,
                timeout=config.OLLAMA_TIMEOUT
            )
            resp.raise_for_status()

            response_json = resp.json()
            logger.info(f"Received response from Ollama: {response_json.keys()}")

            answer = extract_ollama_response(response_json)

            if answer and answer != "خطا در استخراج پاسخ از مدل":
                # پردازش نهایی پاسخ
                answer = post_process_answer(answer)
                logger.info("Successfully extracted and processed answer")
                return answer
            else:
                logger.warning("Could not extract proper answer, retrying...")
                retry_count += 1
                time.sleep(2)

        except requests.exceptions.Timeout:
            logger.error(f"Timeout on attempt {retry_count + 1}")
            retry_count += 1
            if retry_count < max_retries:
                time.sleep(5)
            else:
                return "⏱️ زمان انتظار برای دریافت پاسخ از مدل به پایان رسید. لطفاً دوباره تلاش کنید."

        except requests.exceptions.RequestException as e:
            logger.error(f"Error communicating with Ollama: {str(e)}")
            retry_count += 1
            if retry_count < max_retries:
                time.sleep(3)
            else:
                return f"❌ خطا در ارتباط با Ollama: {str(e)}"

        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}")
            return f"❌ خطای غیرمنتظره: {str(e)}"

    return "متأسفانه نتوانستم پاسخ مناسبی از مدل دریافت کنم."


def save_to_word(query: str, paragraphs_with_scores: List[Tuple[str, float]],
                 answer: str, filename: Optional[str] = None) -> str:
    """ذخیره نتایج در فایل Word با فرمت بهبود یافته"""
    doc = Document()

    # تنظیم فونت پیش‌فرض برای فارسی
    doc.styles['Normal'].font.name = config.DEFAULT_FONT
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT)

    # عنوان اصلی
    title = doc.add_heading('گزارش پرسش و پاسخ', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = config.DEFAULT_FONT
    title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT)
    title.runs[0].font.size = Pt(config.TITLE_FONT_SIZE)

    # تاریخ و زمان
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f'تاریخ: {datetime.now().strftime("%Y/%m/%d - %H:%M")}')
    date_run.font.size = Pt(config.DATE_FONT_SIZE)
    date_run.font.color.rgb = RGBColor(*config.DATE_COLOR)
    date_run.font.name = config.DEFAULT_FONT
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT)

    doc.add_paragraph()  # خط خالی

    # سوال
    question_heading = doc.add_heading('سوال:', level=2)
    question_heading.runs[0].font.name = config.DEFAULT_FONT
    question_heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT)
    question_heading.runs[0].font.size = Pt(config.HEADING_FONT_SIZE)

    q_para = doc.add_paragraph()
    q_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    q_run = q_para.add_run(query)
    q_run.font.size = Pt(config.HEADING_FONT_SIZE)
    q_run.font.bold = True
    q_run.font.color.rgb = RGBColor(*config.QUESTION_COLOR)
    q_run.font.name = config.DEFAULT_FONT
    q_run._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT)

    doc.add_paragraph()  # خط خالی

    # پاراگراف‌های یافت شده
    para_heading = doc.add_heading('پاراگراف‌های مرتبط:', level=2)
    para_heading.runs[0].font.name = config.DEFAULT_FONT
    para_heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT)
    para_heading.runs[0].font.size = Pt(config.HEADING_FONT_SIZE)

    for i, (para, score) in enumerate(paragraphs_with_scores, 1):
        # عنوان پاراگراف
        sub_heading = doc.add_heading(f'پاراگراف {i} (امتیاز: {score:.4f})', level=3)
        sub_heading.runs[0].font.name = config.DEFAULT_FONT
        sub_heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT)
        sub_heading.runs[0].font.size = Pt(config.BODY_FONT_SIZE)

        # متن پاراگراف
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_run = p.add_run(para)
        p_run.font.size = Pt(config.BODY_FONT_SIZE)
        p_run.font.name = config.DEFAULT_FONT
        p_run._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT)

        doc.add_paragraph()  # خط خالی بین پاراگراف‌ها

    # خط جداکننده
    doc.add_paragraph('─' * 50)

    # پاسخ نهایی
    answer_heading = doc.add_heading('پاسخ نهایی Deepseek:', level=2)
    answer_heading.runs[0].font.name = config.DEFAULT_FONT
    answer_heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT)
    answer_heading.runs[0].font.color.rgb = RGBColor(*config.ANSWER_COLOR)
    answer_heading.runs[0].font.size = Pt(config.HEADING_FONT_SIZE)

    # متن پاسخ با فرمت بهتر
    lines = answer.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()  # خط خالی
            continue

        para = doc.add_paragraph()

        # تشخیص نوع خط و اعمال فرمت مناسب
        if re.match(r'^\d+[\-\.]', line):
            # بخش اصلی - bold و بزرگتر
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run(line)
            run.font.bold = True
            run.font.size = Pt(config.HEADING_FONT_SIZE)
        elif line.strip().startswith(('•', '◦', '▪', '▫')):
            # زیربخش - با تورفتگی
            para.paragraph_format.right_indent = Pt(20)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run(line)
            run.font.size = Pt(config.BODY_FONT_SIZE)
        else:
            # متن عادی
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(line)
            run.font.size = Pt(config.BODY_FONT_SIZE)

        # اعمال فونت فارسی
        run.font.name = config.DEFAULT_FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), config.DEFAULT_FONT)
    # نام فایل
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"QA_Report_{timestamp}.docx"

    # ذخیره فایل
    doc.save(filename)
    logger.info(f"Report saved to: {filename}")
    return filename



def main():
    """تابع اصلی برنامه"""
    # حذف فایل‌های ایندکس قدیمی
    if os.path.exists(config.INDEX_BIN_PATH) or os.path.exists(config.META_PKL_PATH):
        user_input = input("⚠️ فایل‌های ایندکس قبلی وجود دارند. آیا می‌خواهید آن‌ها را حذف کنید؟ (بله/خیر): ")
        if user_input.lower() in ['بله', 'yes', 'y']:
            if os.path.exists(config.INDEX_BIN_PATH):
                os.remove(config.INDEX_BIN_PATH)
                logger.info(f"Removed old index file: {config.INDEX_BIN_PATH}")
            if os.path.exists(config.META_PKL_PATH):
                os.remove(config.META_PKL_PATH)
                logger.info(f"Removed old meta file: {config.META_PKL_PATH}")
            print("✅ فایل‌های قدیمی حذف شدند.")

    txt_path = input("📂 مسیر فایل TXT را وارد کنید: ").strip()

    # بررسی پسوند فایل
    if not txt_path.lower().endswith('.txt'):
        print("⚠️ لطفاً یک فایل با پسوند .txt وارد کنید!")
        logger.error("Invalid file type provided")
        exit()

    # ساخت ایندکس
    build_index(txt_path)

    print("\n🤖 سیستم پرسش و پاسخ با Deepseek آماده است!")

    while True:
        query_text = input("\n❓ پرسش خود را وارد کنید (خالی برای خروج): ").strip()
        if not query_text:
            print("👋 خداحافظ!")
            break

        print("\n🔍 در حال جستجو...")
        results = search_by_threshold(query_text,threshold=0.9)

        # نمایش پاراگراف‌های یافت شده
        print("\n📋 پاراگراف‌های مرتبط:")
        top_paragraphs = []
        for i, (para, score) in enumerate(results, 1):
            print(f"\nپاراگراف {i} (امتیاز {score:.4f}):")
            print(para)
            print("-" * 80)
            top_paragraphs.append(para)

        # ارسال به Deepseek
        print("\n🤔 در حال پردازش با Deepseek...")
        answer = ask_deepseek(top_paragraphs, query_text)

        print("\n✅ پاسخ نهایی:")
        print("=" * 80)
        # نمایش خط به خط برای خوانایی بهتر
        for line in answer.split('\n'):
            if line.strip():
                print(line)
        print("=" * 80)



if __name__ == "__main__":
    main()
